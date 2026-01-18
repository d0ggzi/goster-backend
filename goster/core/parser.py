from __future__ import annotations

import re
from collections import defaultdict
from pathlib import Path
from typing import TYPE_CHECKING

from docx import Document
from docx.document import Document as DocxDocument
from docx.table import Table
from docx.text.paragraph import Paragraph

from .model import (
    DocumentModel,
    Element,
    FigureElement,
    FormulaElement,
    HeadingElement,
    ParagraphElement,
    Section,
    STRUCTURAL_KEYWORDS,
    TableElement,
)
from .classifier import ClassificationContext, HeuristicClassifier

if TYPE_CHECKING:
    from .classifier import HeadingClassifier


TABLE_REF_PATTERN = re.compile(r"(?:таблиц[аеуыи]|табл\.)\s*(\d+(?:\.\d+)?)", re.IGNORECASE)
FIGURE_REF_PATTERN = re.compile(r"(?:рисунк[аеуио]|рисунком|рис\.)\s*(\d+(?:\.\d+)?)", re.IGNORECASE)
TABLE_CAPTION_PATTERN = re.compile(r"^Таблица\s+(\d+(?:\.\d+)?)", re.IGNORECASE)
FIGURE_CAPTION_PATTERN = re.compile(r"^Рисунок\s+(\d+(?:\.\d+)?)", re.IGNORECASE)
FORMULA_PATTERN = re.compile(r"[=+\-*/^∫∑∏√≤≥≠∈∉⊂⊃∀∃]\s*.*\((\d+(?:\.\d+)?)\)\s*$")

HEADING_NUMBERED_PATTERN = re.compile(r"^\d+(\.\d+)+\s+[А-ЯA-Z]")
HEADING_SINGLE_NUMBER_PATTERN = re.compile(r"^\d+\s+[А-ЯA-Z]")
CHAPTER_PATTERN = re.compile(r"^ГЛАВА\s+(\d+)", re.IGNORECASE)
NUMBER_PREFIX_PATTERN = re.compile(r"^(\d+(?:\.\d+)*)")


class DocumentParser:
    def __init__(
        self,
        path: Path | str,
        doc: DocxDocument | None = None,
        classifier: "HeadingClassifier | None" = None,
    ):
        self.path = Path(path)
        self.doc: DocxDocument = doc if doc is not None else Document(self.path)
        self.model = DocumentModel()
        self._heading_counters: list[int] = [0, 0, 0, 0, 0]
        self._table_refs: dict[str, list[int]] = defaultdict(list)
        self._figure_refs: dict[str, list[int]] = defaultdict(list)
        self._classifier = classifier or HeuristicClassifier()
        self._parsed_texts: list[str] = []

    @classmethod
    def from_document(
        cls,
        path: Path | str,
        doc: DocxDocument,
        classifier: "HeadingClassifier | None" = None,
    ) -> "DocumentParser":
        return cls(path, doc=doc, classifier=classifier)

    def parse(self) -> DocumentModel:
        self._extract_elements()
        self._build_sections()
        self._resolve_references()
        return self.model

    def _extract_elements(self):
        body = self.doc.element.body
        element_index = 0

        for child in body:
            tag = child.tag.split("}")[-1]

            if tag == "p":
                para = Paragraph(child, self.doc)
                element = self._parse_paragraph(para, element_index)
                self.model.elements.append(element)

                if isinstance(element, HeadingElement):
                    self.model.headings.append(element)
                elif isinstance(element, ParagraphElement):
                    self.model.paragraphs.append(element)
                    self._extract_references(element)

                element_index += 1

            elif tag == "tbl":
                table = Table(child, self.doc)
                table_element = self._parse_table(table, element_index)
                self.model.elements.append(table_element)
                self.model.tables.append(table_element)
                element_index += 1

    def _get_prev_texts(self, count: int = 3) -> list[str]:
        return self._parsed_texts[-count:] if self._parsed_texts else []

    def _parse_paragraph(self, para: Paragraph, index: int) -> Element:
        style_name = para.style.name if para.style else ""
        text = para.text.strip()

        if TABLE_CAPTION_PATTERN.match(text):
            self._parsed_texts.append(text[:100] if text else "")
            return ParagraphElement(index=index, raw=para)

        figure_match = FIGURE_CAPTION_PATTERN.match(text)
        if figure_match:
            figure = FigureElement(index=index, raw=para)
            figure.number = figure_match.group(1)
            figure.caption = text
            self.model.figures.append(figure)
            self._parsed_texts.append(text[:100] if text else "")
            return figure

        formula_match = FORMULA_PATTERN.search(text)
        if formula_match:
            formula = FormulaElement(index=index, raw=para)
            formula.number = formula_match.group(1)
            self.model.formulas.append(formula)
            self._parsed_texts.append(text[:100] if text else "")
            return formula

        ctx = ClassificationContext(
            text=text,
            style_name=style_name,
            prev_texts=self._get_prev_texts(),
        )
        result = self._classifier.classify(ctx)

        if result.element_type == "heading" and result.heading_level is not None:
            if len(text) > 100 and result.confidence < 0.8:
                self._parsed_texts.append(text[:100] if text else "")
                return ParagraphElement(index=index, raw=para)

            heading = HeadingElement(
                index=index,
                raw=para,
                level=result.heading_level,
                classification_source=result.source,
                classification_confidence=result.confidence,
            )

            if not heading.is_structural:
                chapter_match = CHAPTER_PATTERN.match(text.upper())
                if chapter_match:
                    chapter_num = int(chapter_match.group(1))
                    self._heading_counters[0] = chapter_num
                    for i in range(1, 5):
                        self._heading_counters[i] = 0
                    heading.number = str(chapter_num)
                else:
                    heading.number = self._compute_heading_number(result.heading_level)

            self._parsed_texts.append(text[:100] if text else "")
            return heading

        self._parsed_texts.append(text[:100] if text else "")
        return ParagraphElement(index=index, raw=para)

    def _parse_table(self, table: Table, index: int) -> TableElement:
        table_element = TableElement(index=index, raw=table)

        if index > 0:
            prev_element = self.model.elements[index - 1]
            if isinstance(prev_element, ParagraphElement):
                text = prev_element.text.strip()
                match = TABLE_CAPTION_PATTERN.match(text)
                if match:
                    table_element.number = match.group(1)
                    table_element.caption = text
                    table_element.caption_paragraph_index = prev_element.index

        return table_element

    def _compute_heading_number(self, level: int) -> str:
        if level < 1 or level > 5:
            return ""

        self._heading_counters[level - 1] += 1

        for i in range(level, 5):
            self._heading_counters[i] = 0

        parts = [str(self._heading_counters[i]) for i in range(level)]
        return ".".join(parts)

    def _extract_references(self, element: ParagraphElement):
        text = element.text

        for match in TABLE_REF_PATTERN.finditer(text):
            self._table_refs[match.group(1)].append(element.index)

        for match in FIGURE_REF_PATTERN.finditer(text):
            self._figure_refs[match.group(1)].append(element.index)

    def _resolve_references(self):
        for table in self.model.tables:
            if table.number and table.number in self._table_refs:
                table.referenced_at = self._table_refs[table.number]

        for figure in self.model.figures:
            if figure.number and figure.number in self._figure_refs:
                figure.referenced_at = self._figure_refs[figure.number]

    def _build_sections(self):
        section_stack: list[Section] = []
        current_elements: list[Element] = []
        preamble_elements: list[Element] = []
        found_first_heading = False

        for element in self.model.elements:
            if isinstance(element, HeadingElement):
                found_first_heading = True
                if current_elements and section_stack:
                    section_stack[-1].elements.extend(current_elements)
                    for el in current_elements:
                        self.model._element_to_section[el.index] = section_stack[-1]
                current_elements = []

                section = Section(
                    heading=element,
                    level=element.level,
                )

                while section_stack and section_stack[-1].level >= element.level:
                    section_stack.pop()

                if section_stack:
                    section.parent = section_stack[-1]
                    section_stack[-1].children.append(section)
                else:
                    self.model.sections.append(section)

                section_stack.append(section)
                self.model._element_to_section[element.index] = section
            else:
                if found_first_heading:
                    current_elements.append(element)
                else:
                    preamble_elements.append(element)

        if current_elements and section_stack:
            section_stack[-1].elements.extend(current_elements)
            for el in current_elements:
                self.model._element_to_section[el.index] = section_stack[-1]

        self.model.preamble_elements = preamble_elements


def parse_document(
    path: Path | str,
    classifier: "HeadingClassifier | None" = None,
) -> DocumentModel:
    parser = DocumentParser(path, classifier=classifier)
    return parser.parse()
